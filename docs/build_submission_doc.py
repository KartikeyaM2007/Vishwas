from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None
    ImageDraw = None
    ImageFont = None


ROOT = Path(r"E:\Vishwas")
DOCS = ROOT / "docs"
OUT = DOCS / "CityPulse_Project_Description_Kartikeya.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run(run, size=None, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_run(run, size=20 if level == 1 else 15, bold=False, color="000000")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run(run, size=11)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=11)
    return p


def draw_diagram(path, title, boxes, arrows):
    if Image is None:
        return None
    width, height = 1400, 560
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arial.ttf", 34)
        box_font = ImageFont.truetype("arial.ttf", 23)
        small_font = ImageFont.truetype("arial.ttf", 18)
    except Exception:
        title_font = box_font = small_font = ImageFont.load_default()
    draw.text((40, 28), title, fill=(0, 0, 0), font=title_font)
    for box in boxes:
        x, y, w, h, label, sub, fill = box
        draw.rounded_rectangle(
            (x, y, x + w, y + h),
            radius=24,
            fill=fill,
            outline=(44, 82, 130),
            width=3,
        )
        draw.text((x + 22, y + 24), label, fill=(0, 0, 0), font=box_font)
        lines = sub.split("\n")
        for idx, line in enumerate(lines):
            draw.text(
                (x + 22, y + 66 + idx * 26),
                line,
                fill=(60, 60, 60),
                font=small_font,
            )
    for arrow in arrows:
        x1, y1, x2, y2, label = arrow
        draw.line((x1, y1, x2, y2), fill=(20, 93, 180), width=5)
        draw.polygon(
            [(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)],
            fill=(20, 93, 180),
        )
        if label:
            draw.text(
                ((x1 + x2) / 2 - 45, y1 - 30),
                label,
                fill=(20, 93, 180),
                font=small_font,
            )
    image.save(path)
    return path


def add_image(doc, path, caption, width=5.8):
    if not Path(path).exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_run(run, size=9, color="555555")


def build():
    DOCS.mkdir(exist_ok=True)
    architecture = draw_diagram(
        DOCS / "submission_architecture.png",
        "CityPulse System Architecture",
        [
            (45, 135, 260, 155, "Citizen APK", "Flutter app\nVoice, GPS, photo/video\nCommunity feed", (229, 240, 255)),
            (385, 135, 260, 155, "FastAPI Backend", "Assistant planner\nEvidence validation\nReport APIs", (232, 245, 233)),
            (725, 135, 260, 155, "AI + Storage", "Gemini + OpenAI fallback\nCloudinary media\nSupabase DB", (255, 247, 225)),
            (1065, 135, 260, 155, "Admin Dashboard", "Map, review queue\nComments/replies\nStatus actions", (241, 232, 255)),
        ],
        [
            (305, 212, 385, 212, "API"),
            (645, 212, 725, 212, "validate"),
            (985, 212, 1065, 212, "review"),
        ],
    )
    process = draw_diagram(
        DOCS / "submission_process.png",
        "Citizen Report Verification Flow",
        [
            (60, 135, 230, 150, "Describe Issue", "Voice or typed input\nGemini assistant turn", (229, 240, 255)),
            (370, 135, 230, 150, "Capture Context", "GPS location\nPhoto/video proof", (232, 245, 233)),
            (680, 135, 230, 150, "AI Validation", "Gemini vision\nOpenAI fallback\nStrict confidence gate", (255, 247, 225)),
            (990, 135, 250, 150, "Admin Review", "Auto-submit if safe\nManual review if uncertain\nCommunity comments", (241, 232, 255)),
        ],
        [
            (290, 210, 370, 210, ""),
            (600, 210, 680, 210, ""),
            (910, 210, 990, 210, ""),
        ],
    )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("CityPulse: Agentic Civic Issue Reporting Platform")
    set_run(run, size=26, bold=False)
    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Project Description | Author: kartikeya")
    set_run(run, size=11, color="555555")

    add_heading(doc, "Problem Statement Selected")
    add_para(
        doc,
        "Community Hero - Hyperlocal Problem Solver. CityPulse addresses the gap between citizens who observe civic problems and administrators who need reliable, verified, location-aware reports. Traditional complaint systems often accept unclear evidence, lack GPS precision, provide weak status visibility, and give administrators limited tools for prioritizing work.",
    )

    add_heading(doc, "Solution Overview")
    add_para(
        doc,
        "CityPulse combines a Flutter citizen APK, a FastAPI backend, an AI validation layer, Supabase persistence, Cloudinary media storage, and a React admin dashboard. Citizens can describe a civic issue, capture GPS and proof, submit reports, view community issues, and participate through comments. Admins can inspect reports on a dashboard, route issues, review uncertain submissions, and communicate through the community discussion layer.",
    )
    add_image(doc, architecture, "Figure 1. High-level architecture.", width=6.2)
    add_image(doc, process, "Figure 2. Citizen report verification flow.", width=6.2)

    add_heading(doc, "Key Features")
    add_bullets(
        doc,
        [
            "Citizen mobile app with assistant-led reporting, GPS capture, photo/video evidence, issue feed, My Issues, and community comments.",
            "Strict evidence validation workflow using Gemini vision and OpenAI fallback when Gemini is unavailable or rate-limited.",
            "Manual verification fallback for uncertain proof so invalid or low-confidence evidence does not receive automatic rewards.",
            "Admin review queue for manual-review reports, with approve, reject, more-proof, assign, and status update actions.",
            "Reddit-style community discussion for both citizens and admins, including comments and replies on each civic issue.",
            "Clean status labels and color logic for Pending, Manual Review, Approved, Rejected, Resolved, and Needs More Proof.",
            "Backend health/debug endpoints and safe configuration checks that avoid printing secrets.",
            "Leaderboard and reward eligibility logic designed around verified civic contributions.",
        ],
    )

    add_heading(doc, "Screenshots")
    add_image(
        doc,
        DOCS / "civic_doc_home_final.png",
        "Figure 3. Citizen APK home screen with backend-loaded civic reports.",
        width=3.2,
    )
    add_image(
        doc,
        DOCS / "civic_community_feed.png",
        "Figure 4. Citizen APK community feed with report cards and status labels.",
        width=3.2,
    )
    add_image(
        doc,
        DOCS / "admin_dashboard.png",
        "Figure 5. Admin dashboard map and operational sidebar.",
        width=6.2,
    )
    add_image(
        doc,
        DOCS / "admin_review_queue.png",
        "Figure 6. Admin manual review queue.",
        width=6.2,
    )

    add_heading(doc, "Technologies Used")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Layer"
    hdr[1].text = "Technology"
    for cell in hdr:
        set_cell_shading(cell, "F1F3F4")
    rows = [
        ("Mobile", "Flutter, Dart, Riverpod, speech_to_text, flutter_tts, geolocator, image_picker"),
        ("Backend", "FastAPI, Uvicorn, Python, python-dotenv, multipart APIs"),
        ("AI", "Google Gemini, OpenAI vision fallback, prompt-planned assistant turns"),
        ("Storage", "Supabase PostgreSQL, Cloudinary media hosting"),
        ("Admin Frontend", "React, Vite, React Router, Leaflet, Recharts, Framer Motion"),
        ("Tooling", "Android SDK, Gradle, GitHub, local physical Android testing"),
    ]
    for layer, tech in rows:
        cells = table.add_row().cells
        cells[0].text = layer
        cells[1].text = tech

    add_heading(doc, "Google Technologies Utilized")
    add_bullets(
        doc,
        [
            "Google Gemini API for assistant planning, civic issue reasoning, structured report extraction, and vision evidence validation.",
            "Android platform tooling for physical-device APK testing on the iQOO / I2018 phone.",
            "Google Maps/Leaflet-compatible map workflows in the admin dashboard context for geospatial civic operations.",
            "Google Docs target format: this DOCX is prepared so it can be uploaded/imported as the required Google Doc and shared with anyone with the link.",
        ],
    )

    add_heading(doc, "Submission Notes")
    add_para(
        doc,
        "The debug APK was built locally and installed on the connected physical Android phone using API_BASE_URL=http://192.168.18.165:8000. Backend health was verified from the phone after VPN was disabled. The citizen home and community feed load live backend data. The issue detail/comment route had a phone-specific blank-body issue during testing, so it is not presented as a completed screenshot in this submission package.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
